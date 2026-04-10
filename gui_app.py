# -*- coding: utf-8 -*-
"""
蚂蚁安全 - Windows应急排查工具 - 主界面
Copyright (C) 蚂蚁安全 (www.mayisafe.cn)
"""

import os
import sys
import json
import subprocess
import threading
import webbrowser
import tempfile
from datetime import datetime
from pathlib import Path

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext

# 导入检测模块
from emergency_tool import (
    AIModule, MiningDetector, RansomwareDetector, VPNDetector,
    RemoteControlDetector, WebShellDetector, HackerAttackDetector,
    FullSystemScanner
)


# ====== 报告生成器 ======
class ReportGenerator:
    """生成专业调查报告"""
    
    COPYRIGHT = "蚂蚁安全 | www.mayisafe.cn"
    TOOL_NAME = "Windows应急排查工具"
    
    @staticmethod
    def generate(findings, scan_type, host_info, ai_analysis=None):
        """生成HTML报告"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 统计威胁等级
        severe = sum(1 for f in findings if f.get("severity") == "严重")
        high = sum(1 for f in findings if f.get("severity") == "高危")
        medium = sum(1 for f in findings if f.get("severity") == "中危")
        low = sum(1 for f in findings if f.get("severity") == "低危")
        info = sum(1 for f in findings if f.get("severity") == "信息")
        
        # 风险评分
        risk_score = severe * 100 + high * 60 + medium * 30 + low * 10
        if risk_score >= 300:
            risk_level = "极高危"
            risk_color = "#d32f2f"
        elif risk_score >= 150:
            risk_level = "高危"
            risk_color = "#f57c00"
        elif risk_score >= 50:
            risk_level = "中危"
            risk_color = "#fbc02d"
        else:
            risk_level = "低危/安全"
            risk_color = "#388e3c"
        
        # 生成发现项HTML
        findings_html = ""
        for i, finding in enumerate(findings, 1):
            severity = finding.get("severity", "信息")
            sev_color = {
                "严重": "#d32f2f",
                "高危": "#f57c00",
                "中危": "#fbc02d",
                "低危": "#388e3c",
                "信息": "#1976d2"
            }.get(severity, "#757575")
            
            # 显示完整路径
            file_path = finding.get("file", finding.get("path", ""))
            if file_path:
                detail_info = f"路径: {file_path}"
            else:
                detail_info = finding.get("description", "")
            
            findings_html += f"""
            <tr>
                <td>{i}</td>
                <td><span class="severity-badge" style="background:{sev_color}">{severity}</span></td>
                <td>{finding.get('type', '未知')}</td>
                <td>{detail_info}</td>
            </tr>
            """
        
        # AI分析结果
        ai_html = ""
        if ai_analysis:
            ai_html = f"""
            <div class="section">
                <h2>🔍 AI智能分析</h2>
                <div class="ai-content">
                    <div class="ai-item">
                        <strong>威胁等级：</strong>
                        <span class="severity-badge" style="background:{risk_color}">{ai_analysis.get('threat_level', '未知')}</span>
                    </div>
                    <div class="ai-item">
                        <strong>威胁类型：</strong> {ai_analysis.get('threat_type', '未知')}
                    </div>
                    <div class="ai-item">
                        <strong>详细分析：</strong>
                        <p>{ai_analysis.get('analysis', '暂无')}</p>
                    </div>
                    <div class="ai-item">
                        <strong>处置建议：</strong>
                        <ul>
                            {"".join(f"<li>{r}</li>" for r in ai_analysis.get('recommendations', []))}
                        </ul>
                    </div>
                </div>
            </div>
            """
        
        # 先构建HTML模板（不使用f-string，避免{$...}与f-string语法冲突）
        tool_name = ReportGenerator.TOOL_NAME
        copyright = ReportGenerator.COPYRIGHT
        
        html = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>安全应急排查报告 - """ + timestamp + """</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Microsoft YaHei', 'Segoe UI', Arial, sans-serif; background: #f5f5f5; padding: 20px; }
        .container { max-width: 1200px; margin: 0 auto; background: #fff; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .header { background: linear-gradient(135deg, #1976d2, #0d47a1); color: white; padding: 30px; text-align: center; }
        .header h1 { font-size: 28px; margin-bottom: 10px; }
        .header .copyright { font-size: 14px; opacity: 0.9; }
        .summary { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; padding: 30px; }
        .summary-card { background: #f9f9f9; border-radius: 8px; padding: 20px; text-align: center; }
        .summary-card .number { font-size: 36px; font-weight: bold; color: #1976d2; }
        .summary-card .label { color: #666; margin-top: 5px; }
        .risk-banner { background: """ + risk_color + """; color: white; padding: 20px; text-align: center; margin: 20px; border-radius: 8px; }
        .risk-banner .level { font-size: 32px; font-weight: bold; }
        .risk-banner .score { font-size: 14px; opacity: 0.9; margin-top: 5px; }
        .section { padding: 30px; border-bottom: 1px solid #eee; }
        .section h2 { color: #1976d2; margin-bottom: 20px; font-size: 20px; }
        .info-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 15px; }
        .info-item { background: #f9f9f9; padding: 15px; border-radius: 6px; }
        .info-item .label { color: #666; font-size: 12px; }
        .info-item .value { font-size: 16px; margin-top: 5px; word-break: break-all; }
        table { width: 100%; border-collapse: collapse; margin-top: 20px; }
        th, td { padding: 12px; text-align: left; border-bottom: 1px solid #eee; }
        th { background: #f5f5f5; font-weight: 600; }
        tr:hover { background: #f9f9f9; }
        .severity-badge { color: white; padding: 4px 12px; border-radius: 20px; font-size: 12px; font-weight: 600; }
        .ai-content { background: #e3f2fd; padding: 20px; border-radius: 8px; }
        .ai-item { margin-bottom: 15px; }
        .ai-item ul { margin-left: 20px; margin-top: 10px; }
        .ai-item li { margin: 5px 0; }
        .footer { background: #263238; color: #90a4ae; padding: 20px; text-align: center; font-size: 12px; }
        .footer a { color: #64b5f6; text-decoration: none; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🛡️ """ + tool_name + """</h1>
            <p>安全应急排查分析报告</p>
            <p class="copyright">""" + copyright + """</p>
        </div>
        
        <div class="risk-banner">
            <div class="level">""" + risk_level + """</div>
            <div class="score">风险评分: """ + str(risk_score) + """ | 发现威胁: """ + str(severe + high + medium + low) + """ 项</div>
        </div>
        
        <div class="summary">
            <div class="summary-card">
                <div class="number" style="color:#d32f2f">""" + str(severe) + """</div>
                <div class="label">严重威胁</div>
            </div>
            <div class="summary-card">
                <div class="number" style="color:#f57c00">""" + str(high) + """</div>
                <div class="label">高危威胁</div>
            </div>
            <div class="summary-card">
                <div class="number" style="color:#fbc02d">""" + str(medium) + """</div>
                <div class="label">中危风险</div>
            </div>
            <div class="summary-card">
                <div class="number" style="color:#388e3c">""" + str(low + info) + """</div>
                <div class="label">低危/信息</div>
            </div>
        </div>
        
        <div class="section">
            <h2>📋 主机信息</h2>
            <div class="info-grid">
                <div class="info-item">
                    <div class="label">主机名</div>
                    <div class="value">""" + host_info.get('hostname', '未知') + """</div>
                </div>
                <div class="info-item">
                    <div class="label">操作系统</div>
                    <div class="value">""" + host_info.get('os', '未知') + """</div>
                </div>
                <div class="info-item">
                    <div class="label">IP地址</div>
                    <div class="value">""" + host_info.get('ip', '未知') + """</div>
                </div>
                <div class="info-item">
                    <div class="label">当前用户</div>
                    <div class="value">""" + host_info.get('user', '未知') + """</div>
                </div>
                <div class="info-item">
                    <div class="label">扫描类型</div>
                    <div class="value">""" + scan_type + """</div>
                </div>
                <div class="info-item">
                    <div class="label">扫描时间</div>
                    <div class="value">""" + timestamp + """</div>
                </div>
            </div>
        </div>
        
        """ + ai_html + """
        
        <div class="section">
            <h2>🔎 检测结果详情</h2>
            <table>
                <thead>
                    <tr>
                        <th style="width:50px">序号</th>
                        <th style="width:80px">等级</th>
                        <th style="width:100px">类型</th>
                        <th>详情</th>
                    </tr>
                </thead>
                <tbody>
                    """ + (findings_html if findings_html else '<tr><td colspan="4" style="text-align:center">未发现异常</td></tr>') + """
                </tbody>
            </table>
        </div>
        
        <div class="footer">
            <p>""" + copyright + """</p>
            <p>本报告由 """ + tool_name + """ 自动生成 | 报告生成时间: """ + timestamp + """</p>
            <p>免责声明：本报告仅供参考，具体处置请咨询专业人员</p>
        </div>
    </div>
</body>
</html>
        """
        return html
    
    @staticmethod
    def generate_txt(findings, scan_type, host_info):
        """生成文本报告"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        lines = [
            "=" * 70,
            "                   Windows应急排查报告",
            "             蚂蚁安全 | www.mayisafe.cn",
            "=" * 70,
            "",
            f"扫描时间: {timestamp}",
            f"主机名: {host_info.get('hostname', '未知')}",
            f"操作系统: {host_info.get('os', '未知')}",
            f"IP地址: {host_info.get('ip', '未知')}",
            f"当前用户: {host_info.get('user', '未知')}",
            f"扫描类型: {scan_type}",
            "",
            "-" * 70,
            "检测结果",
            "-" * 70,
            ""
        ]
        
        for i, f in enumerate(findings, 1):
            lines.append(f"{i}. [{f.get('severity', '未知')}] {f.get('type', '未知')}")
            # 显示完整路径
            file_path = f.get("file") or f.get("path") or ""
            if file_path:
                lines.append(f"   路径: {file_path}")
            lines.append(f"   {f.get('description', '')}")
            lines.append("")
        
        if not findings:
            lines.append("未发现明显异常。")
            lines.append("")
        
        lines.extend([
            "-" * 70,
            "报告版权: 蚂蚁安全 (www.mayisafe.cn)",
            f"报告生成时间: {timestamp}",
            "=" * 70
        ])
        
        return "\n".join(lines)


# ====== PDF报告生成器 ======
class PDFReportGenerator:
    """生成PDF格式报告"""
    
    @staticmethod
    def generate(findings, scan_type, host_info, ai_analysis=None):
        """生成PDF报告"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            return None, "请先安装reportlab: pip install reportlab"
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 统计威胁等级
        severe = sum(1 for f in findings if f.get("severity") == "严重")
        high = sum(1 for f in findings if f.get("severity") == "高危")
        medium = sum(1 for f in findings if f.get("severity") == "中危")
        low = sum(1 for f in findings if f.get("severity") == "低危")
        info = sum(1 for f in findings if f.get("severity") == "信息")
        
        # 风险评分
        risk_score = severe * 100 + high * 60 + medium * 30 + low * 10
        if risk_score >= 300:
            risk_level = "极高危"
        elif risk_score >= 150:
            risk_level = "高危"
        elif risk_score >= 50:
            risk_level = "中危"
        else:
            risk_level = "低危/安全"
        
        # 创建PDF
        filepath = os.path.join(tempfile.gettempdir(), f"安全排查报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        
        try:
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 中文样式
            try:
                pdfmetrics.registerFont(TTFont('Microsoft YaHei', 'C:/Windows/Fonts/msyh.ttc'))
                chinese_font = 'Microsoft YaHei'
            except:
                chinese_font = 'Helvetica'
            
            title_style = ParagraphStyle('Title', fontName=chinese_font, fontSize=20, alignment=1, spaceAfter=20)
            header_style = ParagraphStyle('Header', fontName=chinese_font, fontSize=14, spaceAfter=10)
            body_style = ParagraphStyle('Body', fontName=chinese_font, fontSize=10, spaceAfter=5)
            
            content = []
            
            # 标题
            content.append(Paragraph("Windows应急排查报告", title_style))
            content.append(Paragraph("蚂蚁安全 | www.mayisafe.cn", ParagraphStyle('Sub', fontName=chinese_font, fontSize=10, alignment=1)))
            content.append(Spacer(1, 20))
            
            # 风险等级
            content.append(Paragraph(f"<b>风险等级:</b> {risk_level} (评分: {risk_score})", body_style))
            content.append(Paragraph(f"<b>威胁统计:</b> 严重:{severe} 高危:{high} 中危:{medium} 低危:{low+info}", body_style))
            content.append(Spacer(1, 20))
            
            # 主机信息
            content.append(Paragraph("主机信息", header_style))
            info_data = [
                ["主机名", host_info.get('hostname', '未知')],
                ["操作系统", host_info.get('os', '未知')],
                ["IP地址", host_info.get('ip', '未知')],
                ["当前用户", host_info.get('user', '未知')],
                ["扫描类型", scan_type],
                ["扫描时间", timestamp],
            ]
            info_table = Table(info_data, colWidths=[100, 400])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ]))
            content.append(info_table)
            content.append(Spacer(1, 20))
            
            # AI分析结果
            if ai_analysis:
                content.append(Paragraph("AI智能分析", header_style))
                content.append(Paragraph(f"<b>威胁等级:</b> {ai_analysis.get('threat_level', '未知')}", body_style))
                content.append(Paragraph(f"<b>威胁类型:</b> {ai_analysis.get('threat_type', '未知')}", body_style))
                content.append(Paragraph(f"<b>分析:</b> {ai_analysis.get('analysis', '暂无')}", body_style))
                recommendations = "<br/>".join([f"{i}. {r}" for i, r in enumerate(ai_analysis.get('recommendations', []), 1)])
                content.append(Paragraph(f"<b>建议:</b><br/>{recommendations}", body_style))
                content.append(Spacer(1, 20))
            
            # 检测结果
            content.append(Paragraph("检测结果详情", header_style))
            
            if findings:
                result_data = [["序号", "等级", "类型", "详情"]]
                for i, f in enumerate(findings[:100], 1):  # 限制100条
                    file_path = f.get("file") or f.get("path") or ""
                    if len(file_path) > 50:
                        file_path = "..." + file_path[-47:]
                    result_data.append([
                        str(i),
                        f.get("severity", ""),
                        f.get("type", ""),
                        file_path or f.get("description", "")[:50]
                    ])
                
                result_table = Table(result_data, colWidths=[40, 60, 80, 320])
                result_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                content.append(result_table)
            else:
                content.append(Paragraph("未发现异常", body_style))
            
            content.append(Spacer(1, 30))
            content.append(Paragraph(f"报告版权: 蚂蚁安全 (www.mayisafe.cn)", ParagraphStyle('Footer', fontName=chinese_font, fontSize=8, alignment=1)))
            content.append(Paragraph(f"生成时间: {timestamp}", ParagraphStyle('Footer', fontName=chinese_font, fontSize=8, alignment=1)))
            
            doc.build(content)
            return filepath, None
            
        except Exception as e:
            return None, str(e)


# ====== Word报告生成器 ======
class WordReportGenerator:
    """生成Word格式报告"""
    
    @staticmethod
    def generate(findings, scan_type, host_info, ai_analysis=None):
        """生成Word报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return None, "请先安装python-docx: pip install python-docx"
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 统计威胁等级
        severe = sum(1 for f in findings if f.get("severity") == "严重")
        high = sum(1 for f in findings if f.get("severity") == "高危")
        medium = sum(1 for f in findings if f.get("severity") == "中危")
        low = sum(1 for f in findings if f.get("severity") == "低危")
        info = sum(1 for f in findings if f.get("severity") == "信息")
        
        # 风险评分
        risk_score = severe * 100 + high * 60 + medium * 30 + low * 10
        if risk_score >= 300:
            risk_level = "极高危"
        elif risk_score >= 150:
            risk_level = "高危"
        elif risk_score >= 50:
            risk_level = "中危"
        else:
            risk_level = "低危/安全"
        
        try:
            doc = Document()
            
            # 标题
            title = doc.add_heading('Windows应急排查报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 版权
            p = doc.add_paragraph('蚂蚁安全 | www.mayisafe.cn')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 风险信息
            doc.add_heading('风险评估', level=1)
            doc.add_paragraph(f'风险等级: {risk_level} (评分: {risk_score})')
            doc.add_paragraph(f'威胁统计: 严重:{severe} 高危:{high} 中危:{medium} 低危:{low+info}')
            
            # 主机信息
            doc.add_heading('主机信息', level=1)
            info_table = doc.add_table(rows=6, cols=2)
            info_table.style = 'Light Grid Accent 1'
            info_data = [
                ("主机名", host_info.get('hostname', '未知')),
                ("操作系统", host_info.get('os', '未知')),
                ("IP地址", host_info.get('ip', '未知')),
                ("当前用户", host_info.get('user', '未知')),
                ("扫描类型", scan_type),
                ("扫描时间", timestamp),
            ]
            for i, (key, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = key
                info_table.rows[i].cells[1].text = value
            
            # AI分析
            if ai_analysis:
                doc.add_heading('AI智能分析', level=1)
                doc.add_paragraph(f"威胁等级: {ai_analysis.get('threat_level', '未知')}")
                doc.add_paragraph(f"威胁类型: {ai_analysis.get('threat_type', '未知')}")
                doc.add_paragraph(f"分析: {ai_analysis.get('analysis', '暂无')}")
                doc.add_paragraph('建议:')
                for r in ai_analysis.get('recommendations', []):
                    doc.add_paragraph(r, style='List Bullet')
            
            # 检测结果
            doc.add_heading('检测结果详情', level=1)
            
            if findings:
                result_table = doc.add_table(rows=min(len(findings)+1, 101), cols=4)
                result_table.style = 'Light Grid Accent 1'
                
                # 表头
                headers = ['序号', '等级', '类型', '详情']
                for i, h in enumerate(headers):
                    result_table.rows[0].cells[i].text = h
                
                # 数据
                for i, f in enumerate(findings[:100]):
                    file_path = f.get("file") or f.get("path") or ""
                    detail = file_path if file_path else f.get("description", "")[:50]
                    result_table.rows[i+1].cells[0].text = str(i+1)
                    result_table.rows[i+1].cells[1].text = f.get("severity", "")
                    result_table.rows[i+1].cells[2].text = f.get("type", "")
                    result_table.rows[i+1].cells[3].text = detail
            else:
                doc.add_paragraph('未发现异常')
            
            # 页脚
            doc.add_paragraph('')
            footer = doc.add_paragraph(f'报告版权: 蚂蚁安全 (www.mayisafe.cn)\n生成时间: {timestamp}')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存
            filepath = os.path.join(tempfile.gettempdir(), f"安全排查报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(filepath)
            return filepath, None
            
        except Exception as e:
            return None, str(e)


# ====== 主窗口类 ======
class EmergencyToolApp:
    """应急排查工具主界面"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("蚂蚁安全 - Windows应急排查工具 v2.0")
        self.root.geometry("1200x850")
        self.root.minsize(1100, 750)
        
        self.ai_module = AIModule()
        self.ai_module.load_config()
        
        self.current_findings = []
        self.current_host_info = {}
        self.scan_type = ""
        self.ai_analysis_result = None
        
        self._setup_styles()
        self._create_menu()
        self._create_ui()
        self._load_host_info()
        
    def _setup_styles(self):
        """设置样式"""
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Title.TLabel", font=("Microsoft YaHei", 24, "bold"), foreground="#1976d2")
        style.configure("Subtitle.TLabel", font=("Microsoft YaHei", 12), foreground="#666")
        style.configure("Copyright.TLabel", font=("Microsoft YaHei", 10), foreground="#999")
        style.configure("ScanButton.TButton", font=("Microsoft YaHei", 11), padding=8)
        
    def _create_menu(self):
        """创建菜单栏"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="导出报告(HTML)", command=lambda: self._export_report("html"))
        file_menu.add_command(label="导出报告(PDF)", command=lambda: self._export_report("pdf"))
        file_menu.add_command(label="导出报告(Word)", command=lambda: self._export_report("docx"))
        file_menu.add_command(label="导出报告(TXT)", command=lambda: self._export_report("txt"))
        file_menu.add_separator()
        file_menu.add_command(label="清空结果", command=self._clear_results)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.root.quit)
        
        # 设置菜单
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="设置", menu=settings_menu)
        settings_menu.add_command(label="AI配置", command=self._show_ai_config)
        settings_menu.add_command(label="保存配置", command=self._save_ai_config)
        
        # 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="帮助", menu=help_menu)
        help_menu.add_command(label="使用说明", command=self._show_help)
        help_menu.add_command(label="关于", command=self._show_about)
        
    def _create_ui(self):
        """创建界面"""
        # 顶部标题
        header_frame = ttk.Frame(self.root)
        header_frame.pack(fill="x", padx=20, pady=15)
        
        title_label = ttk.Label(header_frame, text="🛡️ Windows应急排查工具 v2.0", style="Title.TLabel")
        title_label.pack()
        
        copyright_label = ttk.Label(header_frame, text="蚂蚁安全 | www.mayisafe.cn", style="Copyright.TLabel")
        copyright_label.pack()
        
        # 主机信息卡片
        info_frame = ttk.LabelFrame(self.root, text="主机信息", padding=15)
        info_frame.pack(fill="x", padx=20, pady=(0, 10))
        
        info_inner = ttk.Frame(info_frame)
        info_inner.pack(fill="x")
        
        self.info_labels = {}
        info_items = [
            ("hostname", "主机名:", ""),
            ("os", "操作系统:", ""),
            ("ip", "IP地址:", ""),
            ("user", "当前用户:", ""),
        ]
        
        for i, (key, label, value) in enumerate(info_items):
            row = i // 2
            col = (i % 2) * 4
            ttk.Label(info_inner, text=label, font=("Microsoft YaHei", 10, "bold")).grid(
                row=row, column=col, sticky="w", padx=(0, 5))
            self.info_labels[key] = ttk.Label(info_inner, text=value, foreground="#1976d2")
            self.info_labels[key].grid(row=row, column=col+1, sticky="w", padx=(0, 20))
        
        # 扫描选项卡片
        scan_frame = ttk.LabelFrame(self.root, text="扫描选项", padding=15)
        scan_frame.pack(fill="x", padx=20, pady=(0, 10))
        
        self.scan_var = tk.StringVar(value="full")
        
        # 使用框架分组
        scan_configs = [
            ("full", "完整扫描", "全面检测所有威胁"),
            ("mining", "挖矿检测", "检测挖矿木马"),
            ("ransomware", "勒索检测", "检测勒索病毒"),
            ("vpn", "翻墙检测", "检测VPN/翻墙"),
            ("remote", "远控检测", "检测远控木马"),
            ("webshell", "WebShell", "扫描WebShell"),
            ("hacker", "黑客检测", "综合攻击检测"),
        ]
        
        # 创建网格布局
        for i, (value, text, desc) in enumerate(scan_configs):
            row = i // 4
            col = i % 4
            rb_frame = ttk.Frame(scan_frame)
            rb_frame.grid(row=row, column=col, sticky="w", padx=5, pady=5)
            ttk.Radiobutton(rb_frame, text=text, variable=self.scan_var, value=value).pack(anchor="w")
            ttk.Label(rb_frame, text=desc, foreground="#888", font=("Microsoft YaHei", 8)).pack(anchor="w", padx=(15, 0))
        
        # 按钮区域
        button_frame = ttk.Frame(self.root)
        button_frame.pack(fill="x", padx=20, pady=10)
        
        self.start_btn = ttk.Button(button_frame, text="🚀 开始扫描", style="ScanButton.TButton",
                                    command=self._start_scan)
        self.start_btn.pack(side="left", padx=5)
        
        self.ai_btn = ttk.Button(button_frame, text="🤖 AI分析", style="ScanButton.TButton",
                                  command=self._run_ai_analysis, state="disabled")
        self.ai_btn.pack(side="left", padx=5)
        
        self.export_btn = ttk.Button(button_frame, text="📄 导出报告", style="ScanButton.TButton",
                                      command=lambda: self._export_report("html"), state="disabled")
        self.export_btn.pack(side="left", padx=5)
        
        ttk.Button(button_frame, text="清空", command=self._clear_results).pack(side="left", padx=5)
        
        # 进度条
        self.progress_frame = ttk.Frame(self.root)
        self.progress_frame.pack(fill="x", padx=20, pady=(0, 10))
        self.progress_frame.pack_forget()
        
        self.progress_label = ttk.Label(self.progress_frame, text="准备中...")
        self.progress_label.pack(anchor="w")
        
        self.progress = ttk.Progressbar(self.progress_frame, mode="determinate", length=800)
        self.progress.pack(fill="x", pady=5)
        
        # 结果显示区域
        result_frame = ttk.LabelFrame(self.root, text="检测结果", padding=15)
        result_frame.pack(fill="both", expand=True, padx=20, pady=(0, 10))
        
        # 结果统计
        self.stats_label = ttk.Label(result_frame, text="未扫描", foreground="#888")
        self.stats_label.pack(anchor="w", pady=(0, 10))
        
        # 创建结果树形视图
        tree_frame = ttk.Frame(result_frame)
        tree_frame.pack(fill="both", expand=True)
        
        columns = ("severity", "type", "description")
        self.result_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=12)
        
        self.result_tree.heading("severity", text="威胁等级")
        self.result_tree.heading("type", text="类型")
        self.result_tree.heading("description", text="详情（含完整路径）")
        
        self.result_tree.column("severity", width=100)
        self.result_tree.column("type", width=120)
        self.result_tree.column("description", width=850)
        
        # 滚动条
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.result_tree.yview)
        self.result_tree.configure(yscrollcommand=scrollbar.set)
        
        self.result_tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # 绑定双击事件
        self.result_tree.bind("<Double-1>", self._on_result_double_click)
        
        # AI分析结果显示区
        self.ai_frame = ttk.LabelFrame(self.root, text="AI智能分析", padding=15)
        self.ai_frame.pack(fill="x", padx=20, pady=(0, 10))
        self.ai_frame.pack_forget()
        
        self.ai_text = scrolledtext.ScrolledText(self.ai_frame, height=8, font=("Microsoft YaHei", 10),
                                                  wrap="word")
        self.ai_text.pack(fill="x")
        
        # 状态栏
        self.status_label = tk.Label(self.root, text="就绪", bd=1, relief="sunken", anchor="w",
                                      font=("Microsoft YaHei", 9), bg="#f5f5f5")
        self.status_label.pack(side="bottom", fill="x")
        
    def _load_host_info(self):
        """加载主机信息"""
        try:
            hostname = socket.gethostname()
            
            try:
                result = subprocess.run(
                    ["powershell", "-Command",
                     "(Get-NetIPAddress -AddressFamily IPv4 | Where-Object {$_.InterfaceAlias -notlike '*Loopback*'})[0].IPAddress"],
                    capture_output=True, text=True, timeout=10
                )
                ip = result.stdout.strip() or "获取失败"
            except:
                ip = "获取失败"
            
            try:
                result = subprocess.run(
                    ["powershell", "-Command", "[System.Environment]::UserName"],
                    capture_output=True, text=True, timeout=10
                )
                user = result.stdout.strip() or "获取失败"
            except:
                user = "获取失败"
            
            try:
                result = subprocess.run(
                    ["powershell", "-Command",
                     "(Get-WmiObject Win32_OperatingSystem).Caption"],
                    capture_output=True, text=True, timeout=10
                )
                os = result.stdout.strip() or "获取失败"
            except:
                os = "获取失败"
            
            self.current_host_info = {
                "hostname": hostname,
                "ip": ip,
                "user": user,
                "os": os
            }
            
            self.info_labels["hostname"].config(text=hostname)
            self.info_labels["ip"].config(text=ip)
            self.info_labels["user"].config(text=user)
            self.info_labels["os"].config(text=os)
            
        except Exception as e:
            self.status_label.config(text=f"获取主机信息失败: {str(e)}")
    
    def _start_scan(self):
        """开始扫描"""
        scan_type = self.scan_var.get()
        
        self.start_btn.config(state="disabled")
        self.ai_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.progress_frame.pack(fill="x")
        self.ai_frame.pack_forget()
        
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        def scan_thread():
            try:
                if scan_type == "full":
                    scanner = FullSystemScanner()
                    self.scan_type = "完整系统扫描"
                elif scan_type == "mining":
                    scanner = MiningDetector()
                    self.scan_type = "挖矿木马检测"
                elif scan_type == "ransomware":
                    scanner = RansomwareDetector()
                    self.scan_type = "勒索病毒检测"
                elif scan_type == "vpn":
                    scanner = VPNDetector()
                    self.scan_type = "翻墙/VPN检测"
                elif scan_type == "remote":
                    scanner = RemoteControlDetector()
                    self.scan_type = "远控木马检测"
                elif scan_type == "webshell":
                    scanner = WebShellDetector()
                    self.scan_type = "WebShell检测"
                elif scan_type == "hacker":
                    scanner = HackerAttackDetector()
                    self.scan_type = "黑客攻击检测"
                else:
                    scanner = FullSystemScanner()
                    self.scan_type = "快速扫描"
                
                def progress_callback(value, msg):
                    self.root.after(0, lambda: self._update_progress(value, msg))
                
                findings = scanner.check(progress_callback)
                self.current_findings = findings
                
                self.root.after(0, lambda: self._display_results(findings))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("扫描错误", str(e)))
            finally:
                self.root.after(0, lambda: self._scan_complete())
        
        threading.Thread(target=scan_thread, daemon=True).start()
    
    def _update_progress(self, value, msg):
        """更新进度"""
        self.progress["value"] = value
        self.progress_label.config(text=msg)
        self.status_label.config(text=msg)
    
    def _display_results(self, findings):
        """显示结果"""
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        for f in findings:
            severity = f.get("severity", "信息")
            sev_tag = severity
            
            # 显示完整路径
            file_path = f.get("file") or f.get("path") or ""
            if file_path:
                display_desc = f"{f.get('description', '')} [{file_path}]"
            else:
                display_desc = f.get("description", "")
            
            self.result_tree.insert("", "end", values=(
                severity,
                f.get("type", "未知"),
                display_desc
            ), tags=(sev_tag,))
        
        # 设置颜色标签
        self.result_tree.tag_configure("严重", background="#ffcdd2")
        self.result_tree.tag_configure("高危", background="#ffe0b2")
        self.result_tree.tag_configure("中危", background="#fff9c4")
        self.result_tree.tag_configure("低危", background="#c8e6c9")
        self.result_tree.tag_configure("信息", background="#e3f2fd")
        
        # 更新统计
        severe = sum(1 for f in findings if f.get("severity") == "严重")
        high = sum(1 for f in findings if f.get("severity") == "高危")
        medium = sum(1 for f in findings if f.get("severity") == "中危")
        low = sum(1 for f in findings if f.get("severity") == "低危")
        info_count = sum(1 for f in findings if f.get("severity") == "信息")
        
        total = len(findings)
        self.stats_label.config(
            text=f"扫描完成 | 总计: {total} 项 | 严重: {severe} | 高危: {high} | 中危: {medium} | 低危: {low} | 信息: {info_count}",
            foreground="#1976d2" if total == 0 else ("#d32f2f" if severe > 0 else "#f57c00" if high > 0 else "#388e3c")
        )
        
        self.current_findings = findings
    
    def _scan_complete(self):
        """扫描完成"""
        self.start_btn.config(state="normal")
        self.progress_frame.pack_forget()
        
        if self.current_findings:
            self.ai_btn.config(state="normal")
            self.export_btn.config(state="normal")
        
        self.status_label.config(text=f"{self.scan_type} 完成")
    
    def _run_ai_analysis(self):
        """运行AI分析"""
        if not self.current_findings:
            messagebox.showwarning("提示", "没有可分析的结果")
            return
        
        if not self.ai_module.api_key:
            messagebox.showwarning("AI未配置", "请先在 设置 > AI配置 中配置API密钥")
            return
        
        self.ai_btn.config(state="disabled")
        self.status_label.config(text="正在进行AI分析...")
        
        def ai_thread():
            try:
                findings_str = json.dumps(self.current_findings, ensure_ascii=False, indent=2)
                
                prompt = f"""这是一次安全扫描的检测结果，请分析这些发现并给出专业的安全建议：

扫描类型：{self.scan_type}
主机信息：{json.dumps(self.current_host_info, ensure_ascii=False)}

检测结果：
{findings_str}

请以JSON格式返回分析结果，包含：
- threat_level: 威胁等级（严重/高危/中危/低危/安全）
- threat_type: 主要威胁类型
- analysis: 详细分析说明
- recommendations: 处置建议列表（3-5条）
- evidence: 关键证据列表"""

                result, error = self.ai_module.analyze(prompt)
                
                if error:
                    self.root.after(0, lambda: messagebox.showerror("AI分析错误", error))
                    return
                
                self.ai_analysis_result = result
                self.root.after(0, lambda: self._display_ai_result(result))
                
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("AI分析错误", str(e)))
            finally:
                self.root.after(0, lambda: self.ai_btn.config(state="normal"))
                self.root.after(0, lambda: self.status_label.config(text="AI分析完成"))
        
        threading.Thread(target=ai_thread, daemon=True).start()
    
    def _display_ai_result(self, result):
        """显示AI分析结果"""
        self.ai_frame.pack(fill="x", pady=(0, 10))
        
        self.ai_text.delete("1.0", "end")
        
        output = f"""🤖 AI智能分析结果

{'='*50}
威胁等级：{result.get('threat_level', '未知')}
威胁类型：{result.get('threat_type', '未知')}

{'='*50}
📋 详细分析：
{result.get('analysis', '暂无详细分析')}

{'='*50}
💡 处置建议：
"""
        for i, rec in enumerate(result.get('recommendations', []), 1):
            output += f"   {i}. {rec}\n"
        
        output += f"""
{'='*50}
🔍 关键证据：
"""
        for ev in result.get('evidence', []):
            output += f"   • {ev}\n"
        
        self.ai_text.insert("1.0", output)
    
    def _on_result_double_click(self, event):
        """双击结果项"""
        selection = self.result_tree.selection()
        if selection:
            item = self.result_tree.item(selection[0])
            values = item["values"]
            
            detail_win = tk.Toplevel(self.root)
            detail_win.title("结果详情")
            detail_win.geometry("700x450")
            detail_win.transient(self.root)
            
            text = scrolledtext.ScrolledText(detail_win, font=("Microsoft YaHei", 10), wrap="word")
            text.pack(fill="both", expand=True, padx=10, pady=10)
            
            output = f"威胁等级：{values[0]}\n"
            output += f"类型：{values[1]}\n"
            output += f"\n详情：\n{values[2]}\n"
            
            for finding in self.current_findings:
                file_path = finding.get("file") or finding.get("path") or ""
                desc = finding.get("description", "")
                display = f"{desc} [{file_path}]" if file_path else desc
                if display == values[2]:
                    output += "\n--- 完整信息 ---"
                    for key, value in finding.items():
                        if key not in ["severity", "type", "description"]:
                            output += f"\n{key}：{value}"
                    break
            
            text.insert("1.0", output)
            text.config(state="disabled")
    
    def _export_report(self, fmt="html"):
        """导出报告"""
        if not self.current_findings:
            messagebox.showwarning("提示", "没有可导出的结果")
            return
        
        filetypes = {
            "html": [("HTML文件", "*.html"), ("所有文件", "*.*")],
            "pdf": [("PDF文件", "*.pdf"), ("所有文件", "*.*")],
            "docx": [("Word文件", "*.docx"), ("所有文件", "*.*")],
            "txt": [("文本文件", "*.txt"), ("所有文件", "*.*")],
        }
        
        defaultext = {"html": ".html", "pdf": ".pdf", "docx": ".docx", "txt": ".txt"}
        
        filepath = filedialog.asksaveasfilename(
            defaultextension=defaultext.get(fmt, ".html"),
            filetypes=filetypes.get(fmt, [("所有文件", "*.*")]),
            initialfile=f"安全排查报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        )
        
        if not filepath:
            return
        
        try:
            if fmt == "html":
                html = ReportGenerator.generate(
                    self.current_findings,
                    self.scan_type,
                    self.current_host_info,
                    self.ai_analysis_result
                )
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(html)
                messagebox.showinfo("导出成功", f"HTML报告已保存至：\n{filepath}")
                
            elif fmt == "pdf":
                result_path, error = PDFReportGenerator.generate(
                    self.current_findings,
                    self.scan_type,
                    self.current_host_info,
                    self.ai_analysis_result
                )
                if error:
                    messagebox.showerror("导出失败", error)
                    return
                # 复制到用户选择的位置
                import shutil
                shutil.copy(result_path, filepath)
                messagebox.showinfo("导出成功", f"PDF报告已保存至：\n{filepath}")
                
            elif fmt == "docx":
                result_path, error = WordReportGenerator.generate(
                    self.current_findings,
                    self.scan_type,
                    self.current_host_info,
                    self.ai_analysis_result
                )
                if error:
                    messagebox.showerror("导出失败", error)
                    return
                import shutil
                shutil.copy(result_path, filepath)
                messagebox.showinfo("导出成功", f"Word报告已保存至：\n{filepath}")
                
            else:  # txt
                txt = ReportGenerator.generate_txt(
                    self.current_findings,
                    self.scan_type,
                    self.current_host_info
                )
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(txt)
                messagebox.showinfo("导出成功", f"TXT报告已保存至：\n{filepath}")
            
            if messagebox.askyesno("打开报告", "是否立即打开报告？"):
                webbrowser.open(f"file://{filepath}")
                
        except Exception as e:
            messagebox.showerror("导出错误", str(e))
    
    def _clear_results(self):
        """清空结果"""
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        self.current_findings = []
        self.ai_analysis_result = None
        self.stats_label.config(text="未扫描", foreground="#888")
        self.ai_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.ai_frame.pack_forget()
        self.status_label.config(text="就绪")
    
    def _show_ai_config(self):
        """显示AI配置窗口"""
        config_win = tk.Toplevel(self.root)
        config_win.title("AI配置")
        config_win.geometry("550x420")
        config_win.transient(self.root)
        config_win.resizable(False, False)
        
        frame = ttk.Frame(config_win, padding=20)
        frame.pack(fill="both", expand=True)
        
        ttk.Label(frame, text="AI API 配置", font=("Microsoft YaHei", 14, "bold")).pack(anchor="w", pady=(0, 15))
        
        # API类型
        ttk.Label(frame, text="API类型：").pack(anchor="w", pady=(5, 0))
        api_type_var = tk.StringVar(value=self.ai_module.api_type)
        
        api_types = [
            ("minimax", "MiniMax (默认)"),
            ("kimi", "Kimi (Moonshot)"),
            ("deepseek", "DeepSeek"),
            ("qwen", "阿里千问 (通义千问)"),
            ("openai", "OpenAI兼容"),
        ]
        for value, text in api_types:
            ttk.Radiobutton(frame, text=text, variable=api_type_var, value=value).pack(anchor="w")
        
        # API密钥
        ttk.Label(frame, text="API密钥：").pack(anchor="w", pady=(10, 0))
        api_key_entry = ttk.Entry(frame, width=55, show="*")
        api_key_entry.insert(0, self.ai_module.api_key or "")
        api_key_entry.pack(fill="x", pady=(5, 0))
        
        # API地址
        ttk.Label(frame, text="API地址：").pack(anchor="w", pady=(10, 0))
        endpoint_entry = ttk.Entry(frame, width=55)
        endpoint_entry.insert(0, self.ai_module.endpoint or "https://api.minimax.chat/v1")
        endpoint_entry.pack(fill="x", pady=(5, 0))
        
        # 模型
        ttk.Label(frame, text="模型名称：").pack(anchor="w", pady=(10, 0))
        model_entry = ttk.Entry(frame, width=55)
        model_entry.insert(0, self.ai_module.model or "MiniMax-Text-01")
        model_entry.pack(fill="x", pady=(5, 0))
        
        # 提示信息
        ttk.Label(frame, text="提示：DeepSeek默认模型deepseek-chat，千问默认模型qwen-plus", 
                foreground="#888", font=("Microsoft YaHei", 8)).pack(anchor="w", pady=(5, 0))
        
        def save():
            self.ai_module.api_type = api_type_var.get()
            self.ai_module.save_config(
                api_key_entry.get(),
                endpoint_entry.get(),
                model_entry.get(),
                api_type_var.get()
            )
            config_win.destroy()
            messagebox.showinfo("保存成功", "AI配置已保存")
        
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(pady=20)
        ttk.Button(btn_frame, text="保存", command=save).pack(side="left", padx=5)
        ttk.Button(btn_frame, text="取消", command=config_win.destroy).pack(side="left", padx=5)
    
    def _save_ai_config(self):
        """保存AI配置"""
        self._show_ai_config()
    
    def _show_help(self):
        """显示帮助"""
        help_text = """
Windows应急排查工具 v2.0 - 使用说明

支持的检测场景：
1. 完整扫描 - 全面检测所有威胁类型
2. 挖矿检测 - 检测挖矿木马、矿池连接、异常CPU占用
3. 勒索检测 - 检测勒索病毒、加密文件、勒索信
4. 翻墙检测 - 检测VPN、代理、翻墙软件及异常外连
5. 远控检测 - 检测远控木马、后门程序、可疑端口
6. WebShell检测 - 扫描网站目录中的WebShell
7. 黑客攻击检测 - 全面检测日志、账号、弱口令、注册表、远控、病毒等

使用方法：
1. 选择扫描类型
2. 点击"开始扫描"按钮
3. 等待扫描完成（可实时查看进度）
4. 可使用"AI分析"功能获取智能分析
5. 点击"导出报告"选择格式（HTML/PDF/Word/TXT）

AI功能配置：
- 支持 MiniMax、Kimi、DeepSeek、阿里千问、OpenAI兼容
- 需要在设置中配置相应的API密钥

注意事项：
- 部分功能需要管理员权限运行
- WebShell扫描需要网站目录存在
- PDF/Word导出需要安装相应库：pip install reportlab python-docx

技术支持：www.mayisafe.cn
        """
        
        help_win = tk.Toplevel(self.root)
        help_win.title("使用说明")
        help_win.geometry("550x500")
        
        text = scrolledtext.ScrolledText(help_win, font=("Microsoft YaHei", 10), wrap="word")
        text.pack(fill="both", expand=True, padx=10, pady=10)
        text.insert("1.0", help_text)
        text.config(state="disabled")
    
    def _show_about(self):
        """显示关于"""
        about_text = """
Windows应急排查工具 v2.0

 Copyright (C) 蚂蚁安全 (www.mayisafe.cn)

本工具用于Windows系统的安全应急排查，支持
检测挖矿、勒索、翻墙、远控、WebShell等
多种威胁场景，并可生成专业调查报告。

新增功能：
- 支持DeepSeek、阿里千问AI分析
- 支持PDF、Word报告导出
- 检测结果显示完整路径
- 优化黑客攻击检测范围
- 优化翻墙检测功能

使用方法请查看"帮助 > 使用说明"
        """
        messagebox.showinfo("关于", about_text)


# ====== 主程序入口 ======
def main():
    root = tk.Tk()
    app = EmergencyToolApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
